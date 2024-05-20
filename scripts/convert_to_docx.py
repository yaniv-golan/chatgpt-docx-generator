import sys
import logging
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
import re
import argparse
import pyperclip

# Set up logging
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def add_code_style(document):
    styles = document.styles
    if 'Code' not in styles:
        style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.name = 'Courier New'
        font.size = Pt(10)
        font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        paragraph_format = style.paragraph_format
        paragraph_format.left_indent = Pt(18)
        paragraph_format.space_before = Pt(6)
        paragraph_format.space_after = Pt(6)
        paragraph_format.line_spacing = 1.0

def add_markdown_paragraph(document, text):
    # Split text on bold markers ** and create runs
    parts = re.split(r'(\*\*.*?\*\*)', text)
    p = document.add_paragraph()
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)

def add_bulleted_paragraph(document, text, level=0):
    # Create a bulleted list item with the given indentation level
    p = document.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Pt(18 * level)
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)

def format_text_to_docx(input_text, output_file=None, base_heading_level=1):
    document = Document()
    add_code_style(document)
    logger.debug("Starting to format text to .docx")

    # Split input text into lines
    lines = input_text.split('\n')
    in_code_block = False
    code_paragraph = None

    for line in lines:
        stripped_line = line.strip()

        if stripped_line.startswith('#'):
            # Calculate heading level based on the number of # and base_heading_level
            heading_level = stripped_line.count('#') - (base_heading_level - 1)
            document.add_heading(stripped_line.lstrip('# ').strip(), level=heading_level)
            logger.debug(f"Added heading: {line} as level {heading_level}")
        elif stripped_line.startswith('---'):
            document.add_paragraph().add_run().add_break()
            logger.debug("Added paragraph break")
        elif stripped_line.startswith('```xml'):
            in_code_block = True
            code_paragraph = document.add_paragraph()
            code_paragraph.style = 'Code'
            code_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            logger.debug("Starting XML code block")
        elif stripped_line.startswith('```'):
            in_code_block = False
            code_paragraph = None
            logger.debug("Ending code block")
        elif stripped_line.startswith('- '):
            # Handle bulleted list items
            level = len(line) - len(line.lstrip(' '))
            add_bulleted_paragraph(document, stripped_line[2:], level // 2)
            logger.debug(f"Added bulleted list item: {line}")
        else:
            if in_code_block:
                code_paragraph.add_run(line + '\n')
                logger.debug(f"Added code line: {line}")
            else:
                if re.search(r'\*\*.*?\*\*', line):
                    add_markdown_paragraph(document, line)
                    logger.debug(f"Added markdown paragraph: {line}")
                else:
                    document.add_paragraph(line)
                    logger.debug(f"Added paragraph: {line}")

    if output_file:
        # Save the formatted document
        document.save(output_file)
        logger.info(f"Document saved to {output_file}")
    else:
        # Save the formatted document to a temporary file
        temp_file = 'temp_output.docx'
        document.save(temp_file)
        logger.info(f"Document saved to temporary file {temp_file}")
        return temp_file

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='Convert formatted text to .docx')
    parser.add_argument('--input-file', help='The input text file')
    parser.add_argument('--output-file', help='The output .docx file')
    parser.add_argument('--base-heading-level', type=int, default=1, help='The base heading level (default: 1)')
    parser.add_argument('--input-clipboard', action='store_true', help='Take input from the clipboard')
    parser.add_argument('--output-clipboard', action='store_true', help='Put output in the clipboard')
    parser.add_argument('--debug', action='store_true', help='Enable debug logging')

    args = parser.parse_args()

    if args.debug:
        logger.setLevel(logging.DEBUG)

    if args.input_clipboard:
        input_text = pyperclip.paste()
        logger.info("Input text taken from clipboard")
    else:
        if not args.input_file:
            logger.error("Either --input-file or --input-clipboard must be provided")
            sys.exit(1)
        with open(args.input_file, 'r') as file:
            input_text = file.read()
        logger.info(f"Input text read from file {args.input_file}")

    base_heading_level = args.base_heading_level
    logger.debug(f"Base heading level set to {base_heading_level}")

    output_file = args.output_file if not args.output_clipboard else None

    result_file = format_text_to_docx(input_text, output_file, base_heading_level)

    if args.output_clipboard:
        with open(result_file, 'rb') as file:
            file_data = file.read()
        pyperclip.copy(file_data.decode('latin1'))
        logger.info("Output copied to clipboard")
    else:
        logger.info(f"Output saved to {args.output_file}")
